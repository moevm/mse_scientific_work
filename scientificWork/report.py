# -*- coding: utf-8 -*-
import xlwt
from datetime import date
from openpyxl import Workbook
from docx import Document
from docx.shared import Inches
from scientificWork.models import Publication, UserProfile

font0 = xlwt.Font()
font0.name = 'Times New Roman'
font0.colour_index = 2
font0.bold = True

style0 = xlwt.XFStyle()
style0.font = font0


def print_list_publications_xlsx(c, filters):
    wb = Workbook()
    ws = wb.active
    if len(filters) > 0:
        ws['A1']=u'Применённые фильтры:'
    s = ""
    for x in filters:
        x1 = c[0]
        s += x
        if isinstance(x1[x], unicode):
            s += ' = ' + x1[x]
        else:
            s += ' = ' + x1[x].__str__()
        s+=";  "
    ws['D1']=s

    ws['A2'] = u"Публикация"
    ws['C2'] = u"Тип"
    ws['E2'] = u"Автор"
    ws['G2'] = u"Дата издания"

    i = 4
    for x in c:
        ws['A' + i.__str__()] = x["bookName"]
        ws['C' + i.__str__()] = x["type"]
        ws['E' + i.__str__()] = x["author"]
        ws['G' + i.__str__()] = x["date"].__str__()
        i += 1
    wb.save('scientificWork/static/Publications.xlsx')


def print_list_staff_xlsx(c, filters):
    wb = Workbook()
    ws = wb.active
    if len(filters) > 0:
        ws['A1']=u'Применённые фильтры:'
    s = ""
    for x in filters:
        x1 = c[0]
        s += x
        if isinstance(x1[x], unicode):
            s += ' = ' + x1[x]
        else:
            s += ' = ' + x1[x].__str__()
        s+=";  "
    ws['D1']=s
    ws['A2'] = u'Сотрудник'
    ws['C2'] = u'Тип'
    ws['E2'] = u'Уч.степень'
    ws['G2'] = u'Уч.звание'
    ws['I2'] = u'Дата контракта'

    i = 4
    for x in c:
        ws['A' + i.__str__()] = x["name"]
        ws['C' + i.__str__()] = x["type"]
        ws['E' + i.__str__()] = x["academic_degree"]
        ws['G' + i.__str__()] = x["academic_status"]
        ws['I' + i.__str__()] = x["contract_date"].__str__()
        i += 1
    wb.save('scientificWork/static/Staff.xlsx')


def print_peport_publications_docx(c, filters):
    document = Document()
    document.add_heading(u'Публикации')
    if len(filters) > 0:
        document.add_paragraph(u'Применённые фильтры:')
    str=""
    for x in filters:
        x1 = c[0]
        str+=x
        if isinstance(x1[x], unicode):
            str += ' = ' + x1[x]
        else:
            str += ' = ' + x1[x].__str__()
        str+=";  "
    document.add_paragraph(str)

    table = document.add_table(rows=len(c) + 1, cols=4)
    table.style = 'LightGrid'
    row = table.rows[0]
    row.cells[0].text = u"Название "
    row.cells[1].text = u"Тип "
    row.cells[2].text = u"Автор "
    row.cells[3].text = u"Дата издания"
    i = 1
    for x in c:
        row = table.rows[i]
        row.cells[0].text = x["bookName"]
        row.cells[1].text = x["type"]
        row.cells[2].text = x["author"]
        row.cells[3].text = x["date"].__str__()
        i += 1
    document.save("scientificWork/static/Publications.docx")


def print_peport_staff_docx(c, filters):
    document = Document()
    document.add_heading(u'Сотрудники')
    if len(filters) > 0:
        document.add_paragraph(u'Применённые фильтры:')
    str=""
    for x in filters:
        x1 = c[0]
        str+=x
        if isinstance(x1[x], unicode):
            str += ' = ' + x1[x]
        else:
            str += ' = ' + x1[x].__str__()
        str+=";  "
    document.add_paragraph(str)
    table = document.add_table(rows=len(c) + 1, cols=5)
    table.style = 'LightGrid'
    row = table.rows[0]
    row.cells[0].text = u"Сотрудник "
    row.cells[1].text = u"Тип "
    row.cells[2].text = u"Уч.степень "
    row.cells[3].text = u"Уч.звание"
    row.cells[4].text = u"Дата контракта"
    i=1
    for x in c:
        row = table.rows[i]
        row.cells[0].text = x["name"]
        row.cells[1].text = x["type"]
        row.cells[2].text = x["academic_degree"]
        row.cells[3].text = x["academic_status"]
        row.cells[4].text = x["contract_date"].__str__()
        i += 1
    document.save("scientificWork/static/Staff.docx")
